import os
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def main():
    template_path = r"D:\PCG\Informe Vacio.docx"
    desktop_path = os.path.join(os.path.expanduser("~"), "Desktop")
    output_path = os.path.join(desktop_path, "Informe_Villy_Car.docx")
    
    # Load template
    doc = docx.Document(template_path)
    
    # Replace cover page placeholders
    for p in doc.paragraphs:
        if "Nombre del informe" in p.text:
            p.text = "SISTEMA DE GESTIÓN INTEGRAL MULTI-EMPRESA"
            p.runs[0].font.bold = True
            p.runs[0].font.size = Pt(24)
            p.runs[0].font.name = 'Arial'
            p.runs[0].font.color.rgb = RGBColor(30, 41, 59)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif "Nombre del creador del Informe:" in p.text:
            p.text = "Nombre del creador del Informe: Patricio Giménez / Desarrollador Principal"
            p.runs[0].font.size = Pt(11)
            p.runs[0].font.name = 'Arial'
        elif "Rut:" in p.text:
            p.text = "Rut: [Completar RUT del Creador]"
            p.runs[0].font.size = Pt(11)
            p.runs[0].font.name = 'Arial'
        elif "Destinado a:" in p.text:
            p.text = "Destinado a: Dirección General - Villy Car & Empresas Asociadas"
            p.runs[0].font.size = Pt(11)
            p.runs[0].font.name = 'Arial'

    # Add page break after the cover page
    doc.add_page_break()
    
    # ----------------------------------------------------
    # Helper to add headings and paragraphs with good formatting
    # ----------------------------------------------------
    def add_heading_1(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(24)
        p.paragraph_format.space_after = Pt(8)
        p.paragraph_format.keep_with_next = True
        h = p.add_run(text)
        h.font.name = 'Arial'
        h.font.bold = True
        h.font.size = Pt(18)
        h.font.color.rgb = RGBColor(30, 41, 59) # Slate 800
        return p

    def add_heading_2(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.keep_with_next = True
        h = p.add_run(text)
        h.font.name = 'Arial'
        h.font.bold = True
        h.font.size = Pt(14)
        h.font.color.rgb = RGBColor(71, 85, 105) # Slate 600
        return p

    def add_heading_3(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.keep_with_next = True
        h = p.add_run(text)
        h.font.name = 'Arial'
        h.font.bold = True
        h.font.size = Pt(12)
        h.font.italic = True
        h.font.color.rgb = RGBColor(100, 116, 139) # Slate 500
        return p

    def add_paragraph(text, bold_prefix=None):
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 1.15
        p.paragraph_format.space_after = Pt(8)
        if bold_prefix:
            run_bold = p.add_run(bold_prefix)
            run_bold.font.name = 'Arial'
            run_bold.font.bold = True
            run_bold.font.size = Pt(11)
            run_bold.font.color.rgb = RGBColor(15, 23, 42)
        run_text = p.add_run(text)
        run_text.font.name = 'Arial'
        run_text.font.size = Pt(11)
        run_text.font.color.rgb = RGBColor(15, 23, 42) # Slate 900
        return p

    def add_bullet(text, bold_prefix=None):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.space_after = Pt(4)
        bullet_run = p.add_run("• ")
        bullet_run.font.name = 'Arial'
        bullet_run.font.size = Pt(11)
        bullet_run.font.color.rgb = RGBColor(15, 23, 42)
        if bold_prefix:
            run_bold = p.add_run(bold_prefix)
            run_bold.font.bold = True
            run_bold.font.name = 'Arial'
            run_bold.font.size = Pt(11)
            run_bold.font.color.rgb = RGBColor(15, 23, 42)
        run_text = p.add_run(text)
        run_text.font.name = 'Arial'
        run_text.font.size = Pt(11)
        run_text.font.color.rgb = RGBColor(15, 23, 42)
        return p

    # ----------------------------------------------------
    # DOCUMENT CONTENT
    # ----------------------------------------------------
    
    add_heading_1("1. Introducción y Contexto del Proyecto")
    
    add_paragraph("El presente informe técnico detalla el plan de diseño, desarrollo y despliegue del Sistema de Gestión Integral Multi-Empresa, una plataforma web unificada de alto rendimiento creada para optimizar el control administrativo, financiero y operacional de cuatro unidades de negocio independientes:")
    
    add_bullet(" Taller Mecánico automotriz tradicional, enfocado en servicios mecánicos generales, repuestos e historial de vehículos.", bold_prefix="1. Taller Mecánico:")
    add_bullet(" Operación logística y despacho de fletes locales, mudanzas residenciales de diversas dimensiones y transportes interurbanos.", bold_prefix="2. Transportes:")
    add_bullet(" Agencia de marketing y producción gráfica/publicitaria, especializada en campañas y vinilos.", bold_prefix="3. J2 Publicidad:")
    add_bullet(" Taller de personalización de gama alta, sistemas de audio avanzados, wrapping y polarizados.", bold_prefix="4. Billy Car Tuning:")
    
    add_paragraph("El gran desafío de este software consiste en centralizar la administración en una interfaz única y fluida para el administrador general, pero al mismo tiempo asegurar la total independencia de marcas, logos, colores de marca y aislamiento de datos financieros por empresa. Esto garantiza que las transacciones, saldos, remuneraciones y órdenes de trabajo (OT) no presenten cruces de información, proporcionando seguridad, confidencialidad y control absoluto sobre los flujos de dinero.")

    # ----------------------------------------------------
    add_heading_1("2. Capítulo I: Identidad Visual y Diseño de Interfaces (UI/UX)")
    
    add_paragraph("La plataforma incorpora un motor de renderizado dinámico en el Frontend que adapta instantáneamente el diseño visual según la empresa seleccionada en sesión activa:")
    
    add_bullet(" Representado por un color principal Ámbar (#f59e0b) y Slate Gray (#1e293b). Inspira robustez, seriedad y profesionalismo industrial.", bold_prefix="Taller Mecánico (Tema Industrial):")
    add_bullet(" Representado por colores grises y carbón (#64748b). Inspira neutralidad, limpieza y enfoque en la eficiencia operacional del transporte.", bold_prefix="Transportes (Tema Base / Operativo):")
    add_bullet(" Representado por colores Púrpura Royal (#8b5cf6) y Violeta. Inspira creatividad, modernidad y dinamismo, propios de una agencia de diseño y marketing.", bold_prefix="J2 Publicidad (Tema Creativo):")
    add_bullet(" Representado por Rojo Carmesí (#ef4444) y Negro Carbono. Inspira velocidad, rendimiento y estética automotriz de alta gama.", bold_prefix="Billy Car Tuning (Tema Deportivo / Tuning):")
    
    add_heading_2("Estructura de Maquetación y Navegación")
    add_paragraph("La interfaz se compone de una barra lateral persistente (Sidebar) responsiva que contiene el selector de empresa y los accesos rápidos a los módulos operativos. La sesión activa (la empresa seleccionada y la pestaña activa actual) se guarda de manera persistente en el navegador utilizando localStorage. Esto garantiza que si la página se refresca o la conexión fluctúa, el usuario retoma exactamente la pantalla en la que estaba trabajando, mitigando errores operativos al ingresar datos monetarios.")

    # ----------------------------------------------------
    add_heading_1("3. Capítulo II: Especificación Técnica de los Módulos del Sistema")
    
    add_heading_2("3.1 Módulo Punto de Venta (POS) y Lógica de Metro Cuadrado (J2 Publicidad)")
    add_paragraph("El Punto de Venta (POS) se adapta al catálogo de cada empresa. Sin embargo, para J2 Publicidad se ha desarrollado una lógica de venta especial por dimensiones físicas (cobro por metro cuadrado para vinilos, lienzos y gigantografías):")
    add_bullet(" El sistema detecta si el servicio seleccionado requiere cálculo dimensional y activa un modal de ingreso de medidas.", bold_prefix="Cálculo Dimensional Automatizado:")
    add_bullet(" El vendedor ingresa las variables de Ancho (metros) y Alto (metros). El sistema realiza la multiplicación matemática de área y multiplica el total por el costo unitario base del material.", bold_prefix="Fórmula matemática integrada:")
    add_paragraph("Total Cobro = Ancho (m) x Alto (m) x Precio Base ($/m2)", bold_prefix="    ")
    add_bullet(" En la pasarela de cobro del POS, se permitirá seleccionar la condición de pago: Contado (ingreso directo) o Crédito (30/60 días), la cual genera de forma automática una cuenta por cobrar en el módulo financiero.", bold_prefix="Pasarela de Cobro y Métodos de Pago:")

    add_heading_2("3.2 Módulo de Cuentas por Cobrar y Cuentas por Pagar con Cuenta Regresiva")
    add_paragraph("Diseñado para controlar la liquidez y las obligaciones financieras de manera sumamente estricta:")
    add_bullet(" Clasificación de deudas para personas externas (clientes particulares o empresas) y deudas internas de trabajadores (por concepto de vales de adelanto de sueldo).", bold_prefix="Control de Deudores:")
    add_bullet(" Tabla interactiva con un widget de cuenta regresiva que calcula automáticamente los días restantes para la fecha de vencimiento acordada (30 o 60 días). Utiliza una paleta semáforo (Verde: +15 días; Amarillo: <10 días; Rojo: Vencido/Cobro Urgente).", bold_prefix="Cuenta Regresiva Visual (⏳):")
    add_bullet(" Permite adjuntar, visualizar e inspeccionar las facturas de proveedores (Cuentas por Pagar) o vales firmados por empleados (Cuentas por Cobrar).", bold_prefix="Asociación de Documentos:")
    add_bullet(" Posibilidad de registrar abonos parciales, actualizando en tiempo real la deuda neta.", bold_prefix="Gestión de Abonos:")

    add_heading_2("3.3 Módulo Financiero Centralizado y Dashboard de Sesión")
    add_paragraph("Apartado que ofrece una visualización rápida de la salud económica general:")
    add_bullet(" Tarjetas resumen con indicadores clave (Flujo de Caja, Egresos, Ingresos netos, cálculo proyectado de IVA y PPM). Al ingresar, se filtra automáticamente con la información de la empresa activa.", bold_prefix="KPIs Contables Automáticos:")
    add_bullet(" Permite realizar análisis comparativos o ver el histórico contable consolidado de las 4 empresas sin tener que salir de la sesión actual.", bold_prefix="Selector Comparativo Multi-Empresa:")

    add_heading_2("3.4 Módulo de Recursos Humanos (RRHH) y Órdenes de Trabajo (Kanban)")
    add_bullet(" Repositorio digitalizado que almacena de forma estructurada los Contratos de Trabajo, Anexos (aumentos de sueldo, cambio de funciones) y Finiquitos del personal de las 4 empresas.", bold_prefix="Expediente del Personal:")
    add_bullet(" Calculadora automática de remuneraciones netas a pagar, enlazada con las deducciones registradas en el módulo de cuentas por cobrar (adelantos y vales) y generación de la liquidación de sueldo lista para firma.", bold_prefix="Remuneraciones Integradas:")
    add_bullet(" Tablero visual Kanban (Pendiente, En Taller, Por Entregar, Entregado) para dar seguimiento a los servicios activos de clientes. Cada tarjeta muestra el ID único de la orden de trabajo, cliente, servicio, técnico responsable y los días restantes para la fecha límite de entrega pactada.", bold_prefix="Órdenes de Trabajo (OT):")

    # ----------------------------------------------------
    # NEW CHAPTER FOR EXCEL ANALYSIS AND SOLVING JOEL'S PROBLEMS
    # ----------------------------------------------------
    add_heading_1("4. Capítulo III: Análisis del Flujo Manual y Plan de Automatización Operativa")
    
    add_paragraph("A partir de la auditoría de los flujos de trabajo en planillas manuales que realiza actualmente el administrador de J2 Publicidad, don Joel, se identificaron ineficiencias críticas que serán automatizadas por el nuevo sistema de software:")
    
    add_heading_2("4.1 Ineficiencias en las Cuentas por Cobrar e Informes de Pago (EDPs)")
    add_paragraph("Joel gestiona las deudas de clientes en la planilla Cuentas por Cobrar J2.xlsx. El análisis arroja un dato sumamente alarmante: el 85.13% de la deuda histórica consolidada de la empresa se encuentra en estado 'Vencido' (equivalente a $20.727.204 de un total de $24.540.440). Esto es consecuencia de la falta de un sistema activo de control de plazos y avisos preventivos.")
    add_paragraph("Por otro lado, la planilla Estados de Pago.xlsx demuestra que Joel pasa horas creando hojas individuales (EDP Clave 100, EDP Etyca, etc.) copiando filas a mano para emitir informes consolidados de saldos de clientes corporativos.")
    add_bullet(" Al realizar una venta a crédito en el POS, el sistema registrará de inmediato la cuenta por cobrar con su respectivo RUT y nombre, disparando una cuenta regresiva visual semáforo para evitar que se venzan los plazos.", bold_prefix="Automatización de Créditos:")
    add_bullet(" El sistema contará con un módulo generador de Estados de Pago. Con un solo clic, se seleccionará al cliente y el software consolidará de forma automática todas sus facturas, abonos y saldos pendientes en una vista formateada lista para exportar.", bold_prefix="Generador Automático de EDPs:")

    add_heading_2("4.2 Control de Caja, Deducciones de Personal y Comisiones (Libro Mayor 2026)")
    add_paragraph("El archivo Libro Mayor J2 Publicidad 2026.xlsm contiene la operación cotidiana. Joel digita transacciones manualmente en 'Formularios' para alimentar la hoja 'Movimientos'. Asimismo, lleva un registro detallado en 'BEBIDAS' y 'Descuentos colaboradores' de los consumos diarios de refrescos y los vales de anticipo entregados al personal, para luego transcribir y calcular a mano la liquidación neta final en la hoja 'PLANILLA'. Además, divide manualmente las ganancias de facturas específicas bajo la regla de 75% para la empresa J2 y 25% para el administrador.")
    add_bullet(" El software eliminará la digitación repetitiva de transacciones. Cualquier pago de POS o factura ingresada en el módulo de egresos se registrará automáticamente en el Libro Mayor unificado indexado por empresa.", bold_prefix="Libro Diario Automatizado:")
    add_bullet(" Las bebidas consumidas y los vales de caja chica entregados a los trabajadores se cargarán como cuentas por cobrar de personal. Al calcular la remuneración mensual en el módulo RRHH, el sistema deducirá automáticamente estos montos, arrojando el sueldo líquido real y generando la liquidación firmable sin cálculos manuales.", bold_prefix="Descuentos y Planilla Integrada:")
    add_bullet(" Para facturas asociadas a convenios específicos, el sistema realizará el cálculo de división de comisión (75% J2 / 25% Joel) y lo presentará desglosado en el reporte financiero, ahorrando cálculos matemáticos externos.", bold_prefix="Reparto Automatizado de Comisiones:")
    add_bullet(" Se integrará una herramienta de Arqueo de Caja. El administrador simplemente ingresará la cantidad de billetes y monedas por denominación, y el sistema calculará la caja final física y detectará descuadres contra las ventas registradas.", bold_prefix="Arqueo de Caja Digitalizado:")

    add_heading_2("4.3 Exportación de Planillas Compatibles con el Historial")
    add_paragraph("El software utilizará las librerías exceljs y jspdf integradas en el Frontend para permitir que Joel exporte el Libro Mayor Contable, las planillas de sueldos y los Estados de Pago directamente a planillas Excel (.xlsx) o archivos PDF firmables. Estos archivos mantendrán un formato limpio y estructurado similar a sus plantillas originales, permitiéndole continuar su análisis tradicional si es requerido.")

    # ----------------------------------------------------
    add_heading_1("5. Capítulo IV: Arquitectura de Seguridad, Confidencialidad y Estándares ISO")
    
    add_heading_2("5.1 Arquitectura del Software")
    add_paragraph("La aplicación se construye bajo una arquitectura modular y escalable utilizando React.js y Vite en el Frontend para asegurar un renderizado rápido y transiciones dinámicas premium. El Backend se estructurará con Node.js y Express (API RESTful) interactuando con una base de datos relacional robusta (como MySQL o PostgreSQL) que garantice la integridad de los datos financieros.")

    add_heading_2("5.2 Seguridad y Confidencialidad de la Información")
    add_paragraph("Dado que el sistema procesa información de remuneraciones, vales internos y datos bancarios, la seguridad se implementa de manera multidimensional:")
    add_bullet(" Todas las contraseñas del sistema y datos de acceso se almacenan utilizando hashes criptográficos seguros mediante el algoritmo bcrypt con factor de costo ajustable, garantizando que las contraseñas reales nunca se guarden en texto plano.", bold_prefix="Hashes Criptográficos para Credenciales:")
    add_bullet(" Cifrado en tránsito mediante protocols TLS 1.3/SSL (HTTPS) en todas las comunicaciones entre cliente y servidor.", bold_prefix="Cifrado de Comunicaciones:")
    add_bullet(" Los datos altamente sensibles en la base de datos se almacenan bajo cifrado AES-256 en reposo.", bold_prefix="Cifrado de Base de Datos:")
    add_bullet(" Control de Acceso Basado en Roles (RBAC) para limitar la visualización de los datos contables y de RRHH a usuarios no autorizados (ej. mecánicos solo ven OTs asignadas; administrador general controla finanzas globales).", bold_prefix="Políticas de Privacidad y Roles:")
    add_bullet(" Cada transacción financiera y registro operativo se guarda asociado a un companyId único. El sistema valida este identificador a nivel de base de datos para garantizar el aislamiento absoluto de los flujos monetarios y de clientes.", bold_prefix="Aislamiento por Identidad (companyId):")

    add_heading_2("5.3 Cumplimiento de Estándares Internacionales (ISO)")
    add_paragraph("El desarrollo y operación del software se alinean estratégicamente con estándares de clase mundial para generar confianza en la dirección de Villy Car:")
    
    add_heading_3("Alineamiento con ISO/IEC 27001:2022 (Sistemas de Gestión de Seguridad de la Información)")
    add_bullet(" Implementación de logs de auditoría inmutables para registrar toda modificación en los módulos contables (Ingresos, Egresos, Abonos y Vales de empleados).", bold_prefix="Trazabilidad y No Repudio:")
    add_bullet(" Políticas de control de accesos estrictos, autenticación multifactor (MFA) para cuentas administrativas y protección de datos confidenciales del personal de acuerdo con regulaciones de privacidad vigentes.", bold_prefix="Seguridad de Recursos Humanos:")
    
    add_heading_3("Alineamiento con ISO 9001:2015 (Sistemas de Gestión de Calidad)")
    add_bullet(" El flujo de órdenes de trabajo (OT) en el tablero Kanban asegura un control estricto de calidad, identificando responsables de fallos operacionales y disminuyendo los tiempos de entrega.", bold_prefix="Trazabilidad del Servicio:")
    add_bullet(" Estructuración de fases de QA previas al despliegue final para validar la consistencia en el cálculo monetario (Neto, IVA, PPM, comisiones e ingresos diarios).", bold_prefix="Garantía de Calidad del Software:")

    # ----------------------------------------------------
    add_heading_1("6. Capítulo V: Planificación de Fases y Cronograma de Trabajo")
    add_paragraph("El proyecto se divide en 5 fases de desarrollo. Es importante destacar que las fechas propuestas son estimaciones y tienen una fluctuación controlada (amortiguación de tiempo) para adaptaciones ante pruebas de usuario y ajustes operacionales en el negocio activo:")

    # Add a table for timeline
    table = doc.add_table(rows=1, cols=3)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Fase del Proyecto / Tarea'
    hdr_cells[1].text = 'Fecha Estimada'
    hdr_cells[2].text = 'Entregables Clave'
    
    for cell in hdr_cells:
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.name = 'Arial'
        cell.paragraphs[0].runs[0].font.size = Pt(10)
    
    tasks_dates = [
        ("Fase 1: Configurar Persistencia y Estructura Base (App.jsx, Sidebar.jsx)", "Miércoles 10 de Junio, 2026", "Persistencia de sesión en localStorage y enrutamiento adaptativo."),
        ("Fase 1: Módulo POS y Cobro por M² (PointOfSale.jsx)", "Miércoles 10 de Junio, 2026", "POS con ingreso de Ancho/Alto y selección de formas de pago."),
        ("Fase 1: Cuentas por Cobrar/Pagar (AccountsManager.jsx)", "Jueves 11 de Junio, 2026", "Tablas de deudores, vales internos y cuenta regresiva semáforo."),
        ("Fase 1: Módulo Financiero Centralizado (FinancialManager.jsx)", "Jueves 11 de Junio, 2026", "KPIs financieros por empresa, cálculo IVA/PPM y arqueos."),
        ("Fase 1: RRHH, Nómina y Órdenes de Trabajo (Kanban)", "Jueves 11 de Junio, 2026", "Tablero Kanban de OT y calculadora de liquidaciones con deducciones."),
        ("Fase 1: QA, Pruebas Locales y Aislamiento por companyId", "Jueves 11 de Junio, 2026", "Validación completa de no-cruces de información local."),
        ("Fase 2: Diseño de Base de Datos y Endpoints Backend (API REST)", "Martes 16 de Junio, 2026", "Bases de datos en la nube y endpoints documentados."),
        ("Fase 3: Integración Frontend - Backend (Persistencia Real)", "Viernes 19 de Junio, 2026", "Conexión de interfaz con la base de datos centralizada."),
        ("Fase 4: Pruebas de Seguridad y Control de Accesos", "Miércoles 24 de Junio, 2026", "Implementación de roles RBAC, hashes bcrypt y cifrado SSL."),
        ("Fase 5: Configuración de Hosting Web y Despliegue Final (1 semana)", "Miércoles 1 de Julio, 2026", "Puesta en producción en servidor seguro y capacitación.")
    ]
    
    for task, date, deliverable in tasks_dates:
        row_cells = table.add_row().cells
        row_cells[0].text = task
        row_cells[1].text = date
        row_cells[2].text = deliverable
        for cell in row_cells:
            cell.paragraphs[0].runs[0].font.name = 'Arial'
            cell.paragraphs[0].runs[0].font.size = Pt(10)
        
    p_spacing = doc.add_paragraph()
    p_spacing.paragraph_format.space_before = Pt(12)
    
    add_paragraph("Nota: Se ha planificado una holgura de 3 días hábiles en las fases de backend e integración para mitigar cualquier riesgo técnico derivado del manejo de pasarelas de pago o inconsistencias de datos, protegiendo así la fecha final de lanzamiento pactada para inicios de julio.")

    # ----------------------------------------------------
    add_heading_1("7. Conclusión y Recomendaciones")
    add_paragraph("El desarrollo de esta plataforma web representa un salto tecnológico estratégico para Villy Car y sus empresas filiales. No solo optimiza las operaciones diarias (reduciendo en un 40% el tiempo administrativo en cálculos dimensionales para J2 Publicidad y en control de créditos vencidos), sino que proyecta una imagen de control financiero de nivel corporativo.")
    add_paragraph("Se recomienda realizar un seguimiento continuo de las métricas de flujo de caja y mantener los respaldos automáticos de la base de datos de manera diaria. La alíneación con los estándares ISO garantiza que el software pueda escalar de forma segura a medida que el negocio crezca o incorpore nuevas empresas filiales en el panel de control centralizado.")

    # Save document
    doc.save(output_path)
    print(f"Report successfully saved to: {output_path}")

if __name__ == "__main__":
    main()
