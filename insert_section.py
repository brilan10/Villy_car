import os
import docx
from docx.shared import Pt, RGBColor, Inches

def main():
    desktop_path = os.path.join(os.path.expanduser("~"), "Desktop")
    report_path = os.path.join(desktop_path, "Informe_Villy_Car.docx")
    
    if not os.path.exists(report_path):
        print(f"Error: Report not found at {report_path}")
        return
        
    doc = docx.Document(report_path)
    
    # Let's find the paragraph containing 'Arqueo de Caja Digitalizado'
    target_idx = -1
    for i, p in enumerate(doc.paragraphs):
        if "Arqueo de Caja Digitalizado" in p.text:
            target_idx = i
            break
            
    if target_idx == -1:
        print("Error: Target paragraph 'Arqueo de Caja Digitalizado' not found.")
        return
        
    target_p = doc.paragraphs[target_idx]
    
    # Create the new paragraph at the end first
    new_p = doc.add_paragraph()
    new_p.paragraph_format.left_indent = Inches(0.25)
    new_p.paragraph_format.space_after = Pt(4)
    
    # Populate paragraph
    bullet_run = new_p.add_run("• ")
    bullet_run.font.name = 'Arial'
    bullet_run.font.size = Pt(11)
    bullet_run.font.color.rgb = RGBColor(15, 23, 42)
    
    bold_run = new_p.add_run("Filtro Contable Consolidado (Vista 'Mostrar Todas'): ")
    bold_run.font.bold = True
    bold_run.font.name = 'Arial'
    bold_run.font.size = Pt(11)
    bold_run.font.color.rgb = RGBColor(15, 23, 42)
    
    text_run = new_p.add_run("El Módulo Financiero incorporará un selector inteligente en el encabezado para alternar y consultar los flujos contables de cualquier empresa de manera individual, o bien consolidar la información de las 4 empresas juntas en una sola pantalla unificada con etiquetas dinámicas de color, ahorrando el cambio de sesión constante del administrador.")
    text_run.font.name = 'Arial'
    text_run.font.size = Pt(11)
    text_run.font.color.rgb = RGBColor(15, 23, 42)
    
    # Move new_p._p after target_p._p in the XML tree
    target_p._p.addnext(new_p._p)
    
    # Save the updated document back to the Desktop
    doc.save(report_path)
    print("Successfully added new section to the report on Desktop.")

if __name__ == "__main__":
    main()
